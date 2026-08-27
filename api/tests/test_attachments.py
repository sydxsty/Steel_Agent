# -*- coding: utf-8 -*-
"""临时聊天附件解析和一次性提示词拼接测试。"""

import asyncio
import tempfile
import unittest
from pathlib import Path

from attachment_parser import parse_attachment
from attachment_service import (
    AttachmentManager,
    AttachmentServiceError,
    MAX_FILE_BYTES,
)


class AttachmentParserTests(unittest.TestCase):
    """覆盖无需OCR模型即可稳定执行的文本、Word和Excel格式。"""

    def _prepare_task(self, root: Path, filename: str, writer) -> Path:
        manager = AttachmentManager(root, Path(__file__))
        record = manager.reserve("test-session", filename)
        source = record.task_dir / "source" / record.stored_name
        writer(source)
        manager.finish_upload(record, source.stat().st_size)
        return record.task_dir

    def test_text_to_markdown(self):
        with tempfile.TemporaryDirectory() as temporary:
            task_dir = self._prepare_task(
                Path(temporary),
                "说明.txt",
                lambda path: path.write_text("管线钢附件内容", encoding="utf-8"),
            )
            output = parse_attachment(task_dir)
            markdown = output.read_text(encoding="utf-8")
            self.assertIn("# 附件：说明.txt", markdown)
            self.assertIn("管线钢附件内容", markdown)

    def test_docx_preserves_heading_and_table(self):
        from docx import Document

        def write_docx(path: Path):
            document = Document()
            document.add_heading("工艺要求", level=1)
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "项目"
            table.cell(0, 1).text = "数值"
            table.cell(1, 0).text = "厚度"
            table.cell(1, 1).text = "22 mm"
            document.save(path)

        with tempfile.TemporaryDirectory() as temporary:
            task_dir = self._prepare_task(Path(temporary), "工艺.docx", write_docx)
            markdown = parse_attachment(task_dir).read_text(encoding="utf-8")
            self.assertIn("# 工艺要求", markdown)
            self.assertIn("| 项目 | 数值 |", markdown)
            self.assertIn("| 厚度 | 22 mm |", markdown)

    def test_xlsx_preserves_sheet_and_formula(self):
        from openpyxl import Workbook

        def write_xlsx(path: Path):
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "性能"
            sheet.append(["YS", "目标"])
            sheet.append([560, "=A2+10"])
            workbook.save(path)

        with tempfile.TemporaryDirectory() as temporary:
            task_dir = self._prepare_task(Path(temporary), "性能.xlsx", write_xlsx)
            markdown = parse_attachment(task_dir).read_text(encoding="utf-8")
            self.assertIn("## 工作表：性能", markdown)
            self.assertIn("=A2+10", markdown)


class AttachmentManagerTests(unittest.IsolatedAsyncioTestCase):
    """验证限制、附件区块和发送后立即清理行为。"""

    async def test_prompt_is_current_turn_only_and_truncated(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            manager = AttachmentManager(root, Path(__file__))
            record = manager.reserve("session-a", "long.txt")
            source = record.task_dir / "source" / record.stored_name
            source.write_text("A", encoding="utf-8")
            manager.finish_upload(record, 1)
            (record.task_dir / "attachment.md").write_text(
                "很长的附件内容" * 50_000,
                encoding="utf-8",
            )
            record.status = "ready"

            effective = await manager.build_prompt_and_consume(
                "session-a", [record.attachment_id], "分析附件"
            )
            self.assertIn("【附件内容开始】", effective)
            self.assertIn("后续内容已截断", effective)
            self.assertLessEqual(len(effective) - len("分析附件"), 200_000)
            self.assertNotIn(record.attachment_id, manager.records)
            self.assertFalse(record.task_dir.exists())

    async def test_server_side_file_and_pdf_limits(self):
        with tempfile.TemporaryDirectory() as temporary:
            manager = AttachmentManager(Path(temporary), Path(__file__))
            manager.reserve("session-b", "one.pdf")
            manager.reserve("session-b", "two.pdf")
            with self.assertRaises(AttachmentServiceError):
                manager.reserve("session-b", "three.pdf")
            with self.assertRaises(AttachmentServiceError):
                manager.reserve("session-b", "legacy.doc")
            self.assertEqual(MAX_FILE_BYTES, 10 * 1024 * 1024)


if __name__ == "__main__":
    unittest.main()
